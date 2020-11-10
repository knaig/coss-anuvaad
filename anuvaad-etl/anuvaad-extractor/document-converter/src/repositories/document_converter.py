import requests
import config
import os
import base64
import json
import pandas as pd
from utilities.utils import DocumentUtilities
from docx import Document
from docx.shared import Pt
from docx.shared import Twips, Cm,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Length
from utilities import MODULE_CONTEXT
from anuvaad_auditor.loghandler import log_info, log_exception
import uuid
import xlsxwriter
from jsonpath_rw import jsonpath, parse

class DocumentConversion(object):

    def __init__(self, DOWNLOAD_FOLDER):
        self.DOWNLOAD_FOLDER = DOWNLOAD_FOLDER

    def get_data_from_content_handler(self, record_id, user_id, start_page=0, end_page=0):
        doc_utils = DocumentUtilities()
        try:
            #json_data = {'record_id' : record_id, 'all' : True}
            headers = {"ad-userid" : user_id, "Content-Type": "application/json"}
            request_url = doc_utils.url_generation(config.CONTENT_HANDLER_ENDPOINT, record_id, start_page, end_page)
            log_info("Intiating request to fetch data from %s"%request_url, MODULE_CONTEXT)
            response = requests.get(request_url, headers = headers)
            response_data = response.content
            log_info("Received data from fetch-content end point of content handler", MODULE_CONTEXT)
            dict_str = response_data.decode("UTF-8")
            dict_json_data = json.loads(dict_str)
            return dict_json_data
        except Exception as e:
            log_exception("Can not fetch content in content handler: ", MODULE_CONTEXT, e)

    def convert_page_data_into_dataframes(self, pages):
        try:
            dfs              = []
            page_width       = None
            page_height      = None
            page_layout = {}
            for page in pages:
                text_tops        = []
                text_lefts       = []
                text_widths      = []
                text_heights     = []
                font_sizes       = []
                font_families    = []
                font_colors      = []
                text_values      = []
                b64_images       = []

                if 'images' not in list(page.keys()) or 'text_blocks' not in list(page.keys()):
                    log_info('looks like one of the key is missing {}'.format(page.keys()), MODULE_CONTEXT)
                    continue
                
                images       = page['images']
                texts        = page['text_blocks']
            #     tables       = page['tables']
                page_num     = page['page_no']
                page_width   = page['page_width']
                page_height  = page['page_height']
                page_layout.update({'page_width' : page_width, 'page_height' : page_height})
                for text in texts:
                    text_tops.append(text['text_top'])
                    text_lefts.append(text['text_left'])
                    text_widths.append(text['text_width'])
                    text_heights.append(text['text_height'])
                    font_sizes.append(text['font_size'])
                    font_families.append(text['font_family'])
                    font_colors.append(text['font_color'])
                    b64_images.append(None)
                    
                    text_value = []
                    for processed_text in text['tokenized_sentences']:
                        text_value.append(processed_text['tgt'])        
                    text_values.append(' '.join(text_value))
                for image in images:
                    text_tops.append(image['text_top'])
                    text_lefts.append(image['text_left'])
                    text_widths.append(image['text_width'])
                    text_heights.append(image['text_height'])
                    b64_images.append(image['base64'])
                    text_values.append(None)
                    font_sizes.append(None)
                    font_families.append(None)
                    font_colors.append(None)
                
                df = pd.DataFrame(list(zip(text_tops, text_lefts, text_widths, text_heights,
                                                        text_values, font_sizes, font_families, font_colors, b64_images)), 
                                        columns =['text_top', 'text_left', 'text_width', 'text_height',
                                                    'text', 'font_size', 'font_family', 'font_color', 'base64'])
                df.sort_values('text_top', axis = 0, ascending = True, inplace=True) 
                df = df.reset_index()
                df = df.where(pd.notnull(df), None)
                dfs.append(df)
            return dfs, page_layout
        except Exception as e:
            log_exception("dataframe formation error", MODULE_CONTEXT, e)

    def document_creation(self, dataframes, page_layout, record_id):
        try:
            doc_utils = DocumentUtilities()
            document              = Document()
            section               = document.sections[-1]
            section.orientation   = WD_ORIENT.PORTRAIT
            
            section.page_width    = Cm(doc_utils.get_cms(page_layout['page_width']))
            section.page_height   = Cm(doc_utils.get_cms(page_layout['page_height']))

            section.left_margin   = Cm(1.27)
            section.right_margin  = Cm(1.27)
            section.top_margin    = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            document._body.clear_content()
            
            for index, df in enumerate(dataframes):
                for index, row in df.iterrows():
                    # if row['text'] == None and row['base64'] != None:
                    #     image_path = doc_utils.get_path_from_base64(self.DOWNLOAD_FOLDER, row['base64'])           
                    #     document.add_picture(image_path, width=Cm(doc_utils.get_cms(row['text_width'])), 
                    #                     height=Cm(doc_utils.get_cms(row['text_height'])))
                    #     os.remove(image_path)
                    if row['text'] != None and row['base64'] == None:
                        paragraph                      = document.add_paragraph()
                        paragraph_format               = paragraph.paragraph_format
                        paragraph_format.left_indent   = Cm(doc_utils.get_cms(row['text_left']))
                        if index != df.index[-1] and df.iloc[index + 1]['text_top'] != row['text_top']:
                            pixel = df.iloc[index + 1]['text_top'] - row['text_top'] - row['font_size']
                            #print("pixel", pixel, "next top", df.iloc[index + 1]['text_top'], "top", row['text_top'], "font", row['font_size'])
                            if pixel>0:
                                paragraph_format.space_after = Twips(doc_utils.pixel_to_twips(pixel))
                            else:
                                paragraph_format.space_after = Twips(0)
                        else:
                            paragraph_format.space_after = Twips(0)
                        run                            = paragraph.add_run()
                        if row['font_family'] != None and "Bold" in row['font_family']:
                            run.bold                   = True
                        font                           = run.font
                        font.name                      = 'Arial'
                        if row['font_size'] != None:
                            font.size                      = Twips(doc_utils.pixel_to_twips(row['font_size'])) 
                        run.add_text(row['text'])
                run.add_break(WD_BREAK.PAGE)
            out_filename = os.path.splitext(os.path.basename(record_id.split('|')[0]))[0] + str(uuid.uuid4()) + '_translated.docx'
            output_filepath = os.path.join(self.DOWNLOAD_FOLDER , out_filename)
            document.save(output_filepath)
            return out_filename
        except Exception as e:
            log_exception("dataframe to doc formation failed", MODULE_CONTEXT, e)

    def dummy_doc(self, record_id):
        document              = Document()
        out_filename = os.path.splitext(os.path.basename(record_id))[0] + '_translated.docx'
        output_filepath = os.path.join(self.DOWNLOAD_FOLDER , out_filename)
        document.save(output_filepath)
        return out_filename

    def get_tokenized_sentences(self, json_data):
        try:
            jsonpath_expr = parse('$..tokenized_sentences[*]')
            matches       = jsonpath_expr.find(json_data)
            tokenized_sentences = []
            for match in matches:
                tokenized_sentences.append(match.value)
            return tokenized_sentences
        except Exception as e:
            log_exception("Getting only tokenised sentence object failed", MODULE_CONTEXT, e)

    def generate_xlsx_file(self, record_id, json_data):
        try:
            out_xlsx_filename = os.path.splitext(os.path.basename(record_id.split('|')[0]))[0] + str(uuid.uuid4()) + '.xlsx'
            output_filepath_xlsx = os.path.join(self.DOWNLOAD_FOLDER , out_filename)
            workbook = xlsxwriter.Workbook(output_filepath_xlsx)
            worksheet = workbook.add_worksheet()
            worksheet.write('A1', 'Source Sentence') 
            worksheet.write('B1', 'Target Sentence')
            row, column = 1, 0
            tokenised_sentences = self.get_tokenized_sentences(json_data)
            for tokenised_sentence in tokenised_sentences:
                worksheet.write(row, column, tokenised_sentence['src']) 
                worksheet.write(row, column + 1, tokenised_sentence['tgt']) 
                row += 1
            workbook.close()
            return out_xlsx_filename
        except Exception as e:
            log_exception("xlsx file formation failed", MODULE_CONTEXT, e)